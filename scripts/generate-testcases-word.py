#!/usr/bin/env python3
"""
generate-testcases-word.py
Erzeugt ein Word-Dokument (.docx) aus allen JSON-Testfall-Definitionen.

Verwendung:
    python3 scripts/generate-testcases-word.py [--output PFAD]

Standard-Ausgabedatei: scripts/TSE-Testfaelle.docx
"""

import json
import os
import sys
import argparse
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Pfade
# ---------------------------------------------------------------------------
SCRIPT_DIR = Path(__file__).parent
REPO_ROOT   = SCRIPT_DIR.parent
TC_DIR      = REPO_ROOT / "templates" / "testcases"
INDEX_FILE  = TC_DIR / "_INDEX.json"

# ---------------------------------------------------------------------------
# Farben
# ---------------------------------------------------------------------------
COLOR_HEADER_BG   = RGBColor(0x1F, 0x49, 0x7D)   # BSI-Dunkelblau
COLOR_HEADER_TXT  = RGBColor(0xFF, 0xFF, 0xFF)   # Weiß
COLOR_SUB_BG      = RGBColor(0xBD, 0xD7, 0xEE)   # Hellblau (Schritt-Header)
COLOR_ZEBRA       = RGBColor(0xED, 0xF2, 0xF9)   # Sehr hell für ungerade Zeilen
COLOR_TYPE = {
    "PRECONDITION": RGBColor(0xE2, 0xEF, 0xDA),  # Hellgrün
    "SETUP":        RGBColor(0xFF, 0xF2, 0xCC),  # Hellgelb
    "ACTION":       RGBColor(0xDD, 0xEB, 0xF7),  # Hellblau
    "VERIFY":       RGBColor(0xE2, 0xEF, 0xDA),  # Hellgrün
    "NEGATIVE":     RGBColor(0xFF, 0xE0, 0xCC),  # Hellrot/Orange
    "FP":           RGBColor(0xF2, 0xE2, 0xFF),  # Helllila
}
COLOR_TYPE_LABEL = {
    "PRECONDITION": "Vorbedingung",
    "SETUP":        "Setup",
    "ACTION":       "Aktion",
    "VERIFY":       "Prüfung",
    "NEGATIVE":     "Negativtest",
    "FP":           "Formelle Prüfung",
}

# ---------------------------------------------------------------------------
# Hilfsfunktionen: Tabellen & Zellen
# ---------------------------------------------------------------------------

def set_cell_bg(cell, color: RGBColor):
    """Setzt Hintergrundfarbe einer Tabellenzelle."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), str(val.get("sz", 4)))
            el.set(qn("w:color"), val.get("color", "4472C4"))
            borders.append(el)
    tcPr.append(borders)


def bold_run(para, text, size_pt=None, color=None):
    run = para.add_run(text)
    run.bold = True
    if size_pt:
        run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def add_metadata_table(doc, rows: list[tuple], widths=(4.0, 11.5)):
    """Erstellt eine zweispaltige Metadaten-Tabelle."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, value) in enumerate(rows):
        row = table.rows[i]
        row.cells[0].width = Cm(widths[0])
        row.cells[1].width = Cm(widths[1])
        set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        # Label
        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.color.rgb = COLOR_HEADER_TXT
        r0.font.size = Pt(9)
        # Wert
        p1 = row.cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(str(value) if value is not None else "–")
        r1.font.size = Pt(9)
        if i % 2 == 1:
            set_cell_bg(row.cells[1], COLOR_ZEBRA)
    return table


def add_checks_table(doc, checks: list):
    """Erstellt eine Prüfpunkte-Tabelle."""
    headers = ["#", "Prüfpunkt", "Erwartetes Ergebnis", "Ergebnis"]
    col_widths = [Cm(0.8), Cm(8.5), Cm(4.0), Cm(2.2)]
    table = doc.add_table(rows=1 + len(checks), cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header-Zeile
    hdr_row = table.rows[0]
    for j, h in enumerate(headers):
        cell = hdr_row.cells[j]
        cell.width = col_widths[j]
        set_cell_bg(cell, COLOR_HEADER_BG)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = COLOR_HEADER_TXT
        r.font.size = Pt(8)
    # Datenzeilen
    for i, chk in enumerate(checks):
        row = table.rows[i + 1]
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]
        row.cells[2].width = col_widths[2]
        row.cells[3].width = col_widths[3]
        if i % 2 == 1:
            for c in row.cells:
                set_cell_bg(c, COLOR_ZEBRA)
        vals = [
            str(i + 1),
            chk.get("description") or "",
            chk.get("expected") or "",
            chk.get("result") or "–",
        ]
        for j, v in enumerate(vals):
            p = row.cells[j].paragraphs[0]
            r = p.add_run(v)
            r.font.size = Pt(8)
    return table


def list_to_str(val):
    if isinstance(val, list):
        return ", ".join(str(v) for v in val)
    return str(val) if val else "–"


def dict_to_str(d: dict) -> str:
    if not d:
        return "–"
    parts = []
    for k, v in d.items():
        if isinstance(v, dict):
            inner = ", ".join(f"{ik}: {iv}" for ik, iv in v.items())
            parts.append(f"{k}: {{ {inner} }}")
        else:
            parts.append(f"{k}: {v}")
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Dokument-Aufbau
# ---------------------------------------------------------------------------

def apply_document_styles(doc: Document):
    """Grundlegende Stil-Anpassungen."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    for level, size, bold in [
        ("Heading 1", 14, True),
        ("Heading 2", 11, True),
        ("Heading 3", 10, True),
    ]:
        s = doc.styles[level]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = bold
        s.font.color.rgb = COLOR_HEADER_BG

    # Seitenränder
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)


def add_title_page(doc: Document, index_data: dict):
    """Titelseite."""
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("BSI TSE Testfall-Dokumentation")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = COLOR_HEADER_BG

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_p.add_run(f"Schema-Version {index_data.get('schema_version','3.0.0')}  ·  Stand: {index_data.get('created','')}")
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()

    # Übersichtstabelle
    meta = [
        ("Gesamtanzahl Testfälle",  str(index_data.get("total_testcases",""))),
        ("XML-Quellen",             index_data.get("xml_sources_used","–")),
        ("Regelwerk",               index_data.get("rules_database","–")),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    doc.add_page_break()


def add_toc_placeholder(doc: Document):
    """Fügt einen Inhaltsverzeichnis-Platzhalter ein."""
    doc.add_heading("Inhaltsverzeichnis", level=1)
    p = doc.add_paragraph(
        "Das Inhaltsverzeichnis wird beim Öffnen in Microsoft Word automatisch aktualisiert.\n"
        "(Rechtsklick → Felder aktualisieren)"
    )
    p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    p.runs[0].font.italic = True

    # Word-TOC-Feld einfügen
    paragraph = doc.add_paragraph()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-2" \\h \\z \\u '
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)

    doc.add_page_break()


def render_testcase(doc: Document, tc: dict, tc_index: int):
    """Rendert einen einzelnen Testfall ins Dokument."""
    tc_id    = tc.get("testcase_id", f"TC-{tc_index:03d}")
    title    = tc.get("title", "")
    heading  = f"{tc_id}  –  {title}"

    # ── Überschrift ──────────────────────────────────────────────────────────
    doc.add_heading(heading, level=1)

    # ── Metadaten-Tabelle ────────────────────────────────────────────────────
    variables = tc.get("variables") or {}
    var_str = "\n".join(f"{k}: {v}" for k, v in variables.items()) if variables else "–"
    source_tests = list_to_str(tc.get("source_tests") or [])
    profiles     = list_to_str(tc.get("profiles") or [])

    meta_rows = [
        ("ID",             tc_id),
        ("Modul",          tc.get("module","–")),
        ("Submodul",       tc.get("submodule","–")),
        ("Funktion",       tc.get("function","–")),
        ("Testtyp",        tc.get("test_type","–")),
        ("Priorität",      tc.get("priority","–")),
        ("Profile",        profiles),
        ("Quelltests",     source_tests),
        ("Erstellt",       tc.get("created","–")),
        ("Schema",         tc.get("schema_version","–")),
    ]
    add_metadata_table(doc, meta_rows)
    doc.add_paragraph()

    # ── Beschreibung ─────────────────────────────────────────────────────────
    description = tc.get("description","")
    purpose     = tc.get("purpose","")
    if description:
        p = doc.add_paragraph()
        p.add_run("Beschreibung: ").bold = True
        p.add_run(description).font.size = Pt(9)
    if purpose and purpose.strip():
        p = doc.add_paragraph()
        p.add_run("Zweck: ").bold = True
        p.add_run(purpose.strip()).font.size = Pt(9)

    # ── Variablen ────────────────────────────────────────────────────────────
    if variables:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Variablen:").bold = True
        for var_name, var_desc in variables.items():
            vp = doc.add_paragraph(style="List Bullet")
            vp.add_run(f"{var_name}").bold = True
            vp.add_run(f" – {var_desc}")
            for r in vp.runs:
                r.font.size = Pt(9)

    # ── Schritte ─────────────────────────────────────────────────────────────
    steps = tc.get("steps") or []
    if steps:
        doc.add_paragraph()
        h = doc.add_heading("Testschritte", level=2)

    for step in steps:
        step_num  = step.get("step", "?")
        step_type = step.get("type", "")
        step_title = step.get("title","")
        bg_color   = COLOR_TYPE.get(step_type, RGBColor(0xF0, 0xF0, 0xF0))
        type_label = COLOR_TYPE_LABEL.get(step_type, step_type)

        # Schritt-Überschrift als farbige Tabelle (1 Zeile, 3 Spalten)
        sh_table = doc.add_table(rows=1, cols=3)
        sh_table.style = "Table Grid"
        widths_sh = [Cm(1.2), Cm(2.8), Cm(11.5)]
        row_sh = sh_table.rows[0]
        for c, (txt, w) in enumerate([
            (f"Schritt {step_num}", widths_sh[0]),
            (type_label,            widths_sh[1]),
            (step_title,            widths_sh[2]),
        ]):
            cell = row_sh.cells[c]
            cell.width = w
            set_cell_bg(cell, bg_color)
            p = cell.paragraphs[0]
            r = p.add_run(txt)
            r.bold = (c < 2)
            r.font.size = Pt(9)
        doc.add_paragraph()

        # Schritt-Details
        detail_rows = []
        instr = step.get("instruction","")
        if instr:
            detail_rows.append(("Anweisung", instr))
        func = step.get("function")
        if func:
            detail_rows.append(("Funktion", func))
        user = step.get("user")
        if user:
            detail_rows.append(("Nutzer", user))
        params = step.get("parameters")
        if params:
            detail_rows.append(("Parameter", dict_to_str(params)))
        exp_ret = step.get("expected_return")
        if exp_ret:
            detail_rows.append(("Erwarteter Rückgabewert", dict_to_str(exp_ret)))
        refs = step.get("references")
        if refs:
            detail_rows.append(("Referenzen", list_to_str(refs)))
        src = step.get("source_test")
        if src:
            detail_rows.append(("Quelltest", src))
        notes = step.get("notes")
        if notes:
            detail_rows.append(("Hinweise", notes))

        if detail_rows:
            add_metadata_table(doc, detail_rows, widths=(3.5, 12.0))

        checks = step.get("checks") or []
        if checks:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Prüfpunkte:").bold = True
            p.runs[0].font.size = Pt(9)
            add_checks_table(doc, checks)

        doc.add_paragraph()

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Hauptprogramm
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="TSE Testcase → Word")
    parser.add_argument(
        "--output", "-o",
        default=str(SCRIPT_DIR / "TSE-Testfaelle.docx"),
        help="Ausgabepfad der .docx-Datei",
    )
    parser.add_argument(
        "--filter", "-f",
        default="",
        help="Nur Testfälle deren ID mit diesem Prefix beginnt (z.B. TC-II)",
    )
    args = parser.parse_args()

    # Index laden
    if not INDEX_FILE.exists():
        print(f"[FEHLER] Index-Datei nicht gefunden: {INDEX_FILE}", file=sys.stderr)
        sys.exit(1)
    with open(INDEX_FILE, encoding="utf-8") as f:
        index_data = json.load(f)

    tc_entries = index_data.get("testcases", [])
    if args.filter:
        tc_entries = [e for e in tc_entries if e["id"].startswith(args.filter)]
        print(f"[INFO] Filter '{args.filter}': {len(tc_entries)} Testfälle ausgewählt.")

    print(f"[INFO] Verarbeite {len(tc_entries)} Testfälle ...")

    doc = Document()
    apply_document_styles(doc)
    add_title_page(doc, index_data)
    add_toc_placeholder(doc)

    for i, entry in enumerate(tc_entries, start=1):
        tc_id   = entry["id"]
        tc_file = TC_DIR / f"{tc_id}.json"
        if not tc_file.exists():
            print(f"  [WARNUNG] Datei nicht gefunden: {tc_file}")
            continue
        with open(tc_file, encoding="utf-8") as f:
            tc = json.load(f)
        print(f"  [{i:3d}/{len(tc_entries)}] {tc_id}")
        render_testcase(doc, tc, i)

    out_path = Path(args.output)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"\n[OK] Word-Dokument gespeichert: {out_path}")


if __name__ == "__main__":
    main()
